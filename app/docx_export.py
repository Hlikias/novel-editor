# -*- coding: utf-8 -*-
"""Word(.docx) 格式导出与范本解析（基于 python-docx）。

- DocFormat：导出格式参数（标题 / 正文）
- PRESETS：默认 / 论文 / 散文 预设
- export_docx_formatted(path, text, title, fmt)：按格式生成 docx，Word 一比一呈现
- parse_template(path) -> DocFormat：解析 .doc/.docx 范本的标题/正文格式
- convert_doc_to_docx(path)：.doc 老格式 → .docx（本机 Word/WPS COM 转换）
"""
from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass, field, asdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

FONTS = ["宋体", "黑体", "楷体", "仿宋", "微软雅黑", "华文行楷", "隶书"]
ALIGN_NAMES = {"center": "居中", "left": "左对齐", "right": "右对齐", "justify": "两端对齐"}
ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
ALIGN_REV = {v: k for k, v in ALIGN_MAP.items()}


@dataclass
class DocFormat:
    """导出/范本格式参数（标题 + 正文 + 列表 + 页眉页脚）。"""
    title_font: str = "黑体"
    title_size: float = 16.0            # pt（三号）
    title_bold: bool = True
    title_align: str = "center"         # center/left/right/justify
    body_font: str = "宋体"
    body_size: float = 12.0             # pt（小四）
    first_indent_chars: float = 2.0     # 首行缩进（字符数）
    line_spacing: float = 1.5           # 行距（倍数）
    space_before: float = 0.0           # 段前（pt）
    space_after: float = 6.0            # 段后（pt）
    list_bullet: bool = False           # 范本含无序列表（圆点）
    list_numbered: bool = False         # 范本含有序列表（数字）
    header_text: str = ""               # 页眉文本（空=无）
    footer_text: str = ""               # 页脚文本（空=无）
    name: str = "默认"

    def describe(self) -> str:
        """一行可读描述（用于界面展示）。"""
        parts = [
            f"标题：{self.title_font} {self.title_size:g}磅"
            f"{' 加粗' if self.title_bold else ''} {ALIGN_NAMES.get(self.title_align, '居中')}",
            f"正文：{self.body_font} {self.body_size:g}磅，"
            f"首行缩进 {self.first_indent_chars:g} 字符，"
            f"行距 {self.line_spacing:g} 倍，段前 {self.space_before:g}pt 段后 {self.space_after:g}pt",
        ]
        if self.list_bullet or self.list_numbered:
            kinds = []
            if self.list_bullet:
                kinds.append("圆点")
            if self.list_numbered:
                kinds.append("数字")
            parts.append("列表：" + "、".join(kinds) + "编号")
        if self.header_text:
            parts.append(f"页眉：{self.header_text[:20]}")
        if self.footer_text:
            parts.append(f"页脚：{self.footer_text[:20]}")
        return "；".join(parts)

    def layout_instructions(self) -> str:
        """给 AI 的详细排版说明（导出时程序将按范本一比一套用）。"""
        lines = [
            "【排版规范（导出时程序将按范本一比一套用，你只需输出内容与结构）】",
            f"- 标题：{self.title_font} {self.title_size:g}磅"
            f"{' 加粗' if self.title_bold else ''}，{ALIGN_NAMES.get(self.title_align, '居中')}",
            f"- 正文：{self.body_font} {self.body_size:g}磅，首行缩进 {self.first_indent_chars:g} 字符，"
            f"行距 {self.line_spacing:g} 倍，段前 {self.space_before:g}pt 段后 {self.space_after:g}pt",
        ]
        if self.list_bullet or self.list_numbered:
            kinds = []
            if self.list_bullet:
                kinds.append("无序列表（圆点）")
            if self.list_numbered:
                kinds.append("有序列表（数字编号）")
            lines.append(f"- 列表：文档中包含{'、'.join(kinds)}，请把列表内容逐项分行，"
                         f"无序项行首加「- 」，有序项行首加「1. 」（连续递增）")
        if self.header_text:
            lines.append(f"- 页眉：{self.header_text}")
        if self.footer_text:
            lines.append(f"- 页脚：{self.footer_text}")
        return "\n".join(lines)

    def to_config(self) -> dict:
        return asdict(self)

    @staticmethod
    def from_config(d: dict) -> "DocFormat":
        if not isinstance(d, dict):
            return DocFormat()
        data = {k: v for k, v in d.items() if k in DocFormat.__dataclass_fields__}
        return DocFormat(**data)


# 预设
PRESETS: dict[str, DocFormat] = {
    "默认": DocFormat(name="默认"),
    "论文": DocFormat(title_font="黑体", title_size=16.0, title_bold=True, title_align="center",
                      body_font="宋体", body_size=12.0, first_indent_chars=2.0,
                      line_spacing=1.5, space_before=0.0, space_after=6.0, name="论文"),
    "散文": DocFormat(title_font="黑体", title_size=18.0, title_bold=True, title_align="center",
                      body_font="楷体", body_size=12.0, first_indent_chars=2.0,
                      line_spacing=1.5, space_before=0.0, space_after=4.0, name="散文"),
    "公文": DocFormat(title_font="方正小标宋简体", title_size=22.0, title_bold=False, title_align="center",
                      body_font="仿宋", body_size=16.0, first_indent_chars=2.0,
                      line_spacing=1.0, space_before=0.0, space_after=0.0, name="公文"),
}


def _set_run_font(run, font: str):
    """同时设置西文与中文（eastAsia）字体，保证 Word 中一比一。"""
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def _set_first_line_indent(p, chars: float, body_size_pt: float):
    """首行缩进按「字符」设置：w:firstLineChars（Word 优先按字符计）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.get_or_add_ind()
    ind.set(qn("w:firstLineChars"), str(int(max(0.0, chars) * 100)))
    # 兜底：字符数 × 字号(pt) × 20 twips/pt
    ind.set(qn("w:firstLine"), str(int(max(0.0, chars) * body_size_pt * 20)))


# 列表行识别：无序（- • · *）与有序（1. 1) （1））
_BULLET_RE = None
_NUMBER_RE = None


def _list_kind(line: str) -> str | None:
    """返回 'bullet' / 'number' / None（行是否为列表项）。"""
    global _BULLET_RE, _NUMBER_RE
    import re
    if _BULLET_RE is None:
        _BULLET_RE = re.compile(r"^\s*[-•·*‣]\s+\S")
        _NUMBER_RE = re.compile(r"^\s*(?:\d+[\.、)．]|（\d+）|[一二三四五六七八九十]+[、.．])\s+\S")
    if _BULLET_RE.match(line):
        return "bullet"
    if _NUMBER_RE.match(line):
        return "number"
    return None


def export_docx_formatted(path: str, text: str, title: str = "",
                          fmt: DocFormat | None = None) -> None:
    """按格式生成 docx：标题 + 正文 + 列表 + 页眉页脚，与范本格式一比一。

    列表：正文中以「- 」「1. 」等标记开头的行导出为 Word 原生列表
    （List Bullet / List Number 样式，自动编号与圆点）。
    页眉页脚：fmt.header_text / footer_text 写入节页眉页脚（居中，正文同字体小一号）。"""
    fmt = fmt or DocFormat()
    doc = Document()
    # ---- 页眉页脚 ----
    section = doc.sections[0]
    if fmt.header_text:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(fmt.header_text)
        _set_run_font(hr, fmt.body_font)
        hr.font.size = Pt(max(9.0, fmt.body_size - 2.0))
    if fmt.footer_text:
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(fmt.footer_text)
        _set_run_font(fr, fmt.body_font)
        fr.font.size = Pt(max(9.0, fmt.body_size - 2.0))

    # ---- 标题 ----
    if title:
        p = doc.add_paragraph()
        p.alignment = ALIGN_MAP.get(fmt.title_align, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(title)
        _set_run_font(r, fmt.title_font)
        r.font.size = Pt(fmt.title_size)
        r.font.bold = fmt.title_bold
        p.paragraph_format.space_after = Pt(12)

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    paras: list[str] = []
    for ln in lines:
        if ln.strip():
            paras.append(ln.strip())
        elif paras and paras[-1] != "":   # 保留一个空行分段
            paras.append("")

    # ---- 正文 / 列表 ----
    for ln in paras:
        if not ln:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue
        kind = _list_kind(ln)
        if kind:
            # 列表项：用 Word 原生列表样式，去掉行首标记
            text_body = ln
            if kind == "bullet":
                text_body = text_body.lstrip("-•·*‣ ").lstrip()
                p = doc.add_paragraph(style="List Bullet")
            else:
                text_body = _strip_number_prefix(text_body)
                p = doc.add_paragraph(style="List Number")
            pf = p.paragraph_format
            pf.line_spacing = fmt.line_spacing
            pf.space_after = Pt(fmt.space_after)
            r = p.add_run(text_body)
            _set_run_font(r, fmt.body_font)
            r.font.size = Pt(fmt.body_size)
            continue
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = fmt.line_spacing
        pf.space_before = Pt(fmt.space_before)
        pf.space_after = Pt(fmt.space_after)
        if fmt.first_indent_chars > 0:
            _set_first_line_indent(p, fmt.first_indent_chars, fmt.body_size)
        r = p.add_run(ln)
        _set_run_font(r, fmt.body_font)
        r.font.size = Pt(fmt.body_size)
    doc.save(path)


def _strip_number_prefix(line: str) -> str:
    """去掉有序列表行首编号（1. / 1、 / （1） / 一、 等）。"""
    import re
    return re.sub(r"^\s*(?:\d+[\.、)．]|（\d+）|[一二三四五六七八九十]+[、.．])\s*", "", line).strip()


# ---------------------------------------------------------------- 范本解析

def _align_name(paragraph) -> str:
    a = paragraph.alignment
    return ALIGN_REV.get(a, "justify") if a is not None else "justify"


def _run_font(run) -> str:
    name = run.font.name or ""
    try:
        rPr = run._element.rPr
        if rPr is not None and rPr.rFonts is not None:
            east = rPr.rFonts.get(qn("w:eastAsia")) or ""
            if east:
                return east
    except Exception:  # noqa: BLE001
        pass
    return name or "宋体"


def _run_size(run) -> float:
    try:
        sz = run.font.size
        return float(sz.pt) if sz is not None else 12.0
    except Exception:  # noqa: BLE001
        return 12.0


def _para_first_indent_chars(p, body_size: float) -> float:
    """读取首行缩进（字符数）：firstLineChars 优先，其次 firstLine/EMU 换算。"""
    try:
        pPr = p._p.pPr
        if pPr is not None and pPr.ind is not None:
            flc = pPr.ind.get(qn("w:firstLineChars"))
            if flc:
                return int(flc) / 100.0
            fl = pPr.ind.get(qn("w:firstLine"))
            if fl:
                return (int(fl) / 20.0) / max(1.0, body_size)   # twips→pt→字符
    except Exception:  # noqa: BLE001
        pass
    try:
        fli = p.paragraph_format.first_line_indent
        if fli is not None:
            return (fli.pt or 0.0) / max(1.0, body_size)
    except Exception:  # noqa: BLE001
        pass
    return 0.0


def _para_line_spacing(p) -> float:
    pf = p.paragraph_format
    try:
        ls = pf.line_spacing
        if isinstance(ls, float) and ls > 0:
            return ls
    except Exception:  # noqa: BLE001
        pass
    return 1.0


def _para_space(p, attr: str) -> float:
    try:
        v = getattr(p.paragraph_format, attr)
        return float(v.pt) if v is not None else 0.0
    except Exception:  # noqa: BLE001
        return 0.0


def parse_template(path: str) -> DocFormat:
    """解析范本 .doc/.docx 的标题与正文格式参数。

    取第一个非空段落为标题样式来源，其余段落为正文样式来源。
    .doc 老格式先经本机 Word/WPS 转换为临时 .docx。"""
    docx_path = convert_doc_to_docx(path) if path.lower().endswith(".doc") else path
    try:
        doc = Document(docx_path)
    finally:
        if docx_path != path:
            try:
                os.remove(docx_path)
            except Exception:  # noqa: BLE001
                pass
    paras = [p for p in doc.paragraphs if p.text.strip()]
    if not paras:
        return DocFormat(name="范本")
    title_p = paras[0]
    body_p = paras[1] if len(paras) > 1 else paras[0]

    # ---- 标题 ----
    tr = title_p.runs[0] if title_p.runs else None
    fmt = DocFormat(name="范本")
    if tr is not None:
        fmt.title_font = _run_font(tr)
        fmt.title_size = _run_size(tr)
        fmt.title_bold = bool(tr.font.bold)
    fmt.title_align = _align_name(title_p)

    # ---- 正文 ----
    br = body_p.runs[0] if body_p.runs else None
    if br is not None:
        fmt.body_font = _run_font(br)
        fmt.body_size = _run_size(br)
    fmt.first_indent_chars = _para_first_indent_chars(body_p, fmt.body_size)
    fmt.line_spacing = _para_line_spacing(body_p)
    fmt.space_before = _para_space(body_p, "space_before")
    fmt.space_after = _para_space(body_p, "space_after")

    # ---- 列表检测（样式含 List，或行首为列表标记） ----
    for p in paras[1:]:
        try:
            st = p.style.name or ""
        except Exception:  # noqa: BLE001
            st = ""
        if "List" in st:
            if "Number" in st:
                fmt.list_numbered = True
            else:
                fmt.list_bullet = True
            continue
        kind = _list_kind(p.text)
        if kind == "bullet":
            fmt.list_bullet = True
        elif kind == "number":
            fmt.list_numbered = True

    # ---- 页眉 / 页脚 ----
    try:
        sec = doc.sections[0]
        for part, attr in ((sec.header, "header_text"), (sec.footer, "footer_text")):
            try:
                texts = [pp.text.strip() for pp in part.paragraphs if pp.text.strip()]
            except Exception:  # noqa: BLE001
                texts = []
            if texts:
                setattr(fmt, attr, texts[0])
    except Exception:  # noqa: BLE001
        pass
    return fmt


def convert_doc_to_docx(path: str) -> str:
    """把老式 .doc 用本机 Word/WPS 转为 .docx，返回临时 docx 路径。

    Word 不可用时尝试 WPS（Kwps.Application）；都没有则抛 RuntimeError。"""
    tmp_docx = os.path.join(tempfile.mkdtemp(prefix="novel_doc_"), "converted.docx")
    last_err = None
    for prog in ("Word.Application", "Kwps.Application"):
        try:
            import win32com.client
            app = win32com.client.DispatchEx(prog)
            try:
                doc = app.Documents.Open(os.path.abspath(path), ReadOnly=True)
                try:
                    doc.SaveAs2(tmp_docx, FileFormat=16)   # 16 = wdFormatXMLDocument(.docx)
                except Exception:  # noqa: BLE001
                    doc.SaveAs(tmp_docx, FileFormat=16)
                doc.Close(False)
                return tmp_docx
            finally:
                try:
                    app.Quit()
                except Exception:  # noqa: BLE001
                    pass
        except Exception as e:  # noqa: BLE001
            last_err = e
    raise RuntimeError(
        "无法读取 .doc 老格式：本机未检测到可用的 Word/WPS。\n"
        "请在 Word 中把范本另存为 .docx 后再试。" + (f"（{last_err}）" if last_err else "")
    )
