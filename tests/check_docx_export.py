# -*- coding: utf-8 -*-
"""Word 格式导出与范本解析测试：
1) 按 DocFormat 生成 docx → python-docx 打开验证 标题/正文 格式（一比一）
2) 构造范本 → parse_template 解析参数一致
3) 格式设置弹窗：预设切换与 fmt() 往返
4) 主窗口：记住设置后导出 docx、按范本导出(plain) 流程
"""
import os
import sys
import tempfile

os.environ["QT_QPA_PLATFORM"] = "offscreen"
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from PySide6.QtWidgets import QApplication, QFileDialog

import app.main_window as _mw
_mw.save_config = lambda cfg: None
from app.main_window import MainWindow
from app.docx_export import (
    FONTS, DocFormat, PRESETS, export_docx_formatted, parse_template,
)
from app.dialogs.export_format_dialog import ExportFormatDialog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ok = True


def check(name, cond):
    global ok
    print(f"{'PASS' if cond else 'FAIL'} - {name}")
    if not cond:
        ok = False


app = QApplication(sys.argv)

# ---------- 1) 按格式生成 docx 并回读验证 ----------
fmt = DocFormat(title_font="黑体", title_size=16.0, title_bold=True, title_align="center",
                body_font="宋体", body_size=12.0, first_indent_chars=2.0,
                line_spacing=1.5, space_before=0.0, space_after=6.0)
d = tempfile.mkdtemp()
out = os.path.join(d, "out.docx")
export_docx_formatted(out, "第一段内容。\n第二段内容。", title="我的文章", fmt=fmt)

doc = Document(out)
paras = [p for p in doc.paragraphs if p.text.strip()]
check("生成文档有标题+正文", len(paras) == 3 and paras[0].text == "我的文章")
title_p = paras[0]
tr = title_p.runs[0]
check("标题字体黑体", tr.font.name == "黑体" or (tr._element.rPr.rFonts is not None and tr._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"))
check("标题字号 16pt", tr.font.size == Pt(16.0))
check("标题加粗", tr.font.bold is True)
check("标题居中", title_p.alignment == WD_ALIGN_PARAGRAPH.CENTER)

body_p = paras[1]
br = body_p.runs[0]
check("正文字体宋体", br._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体")
check("正文字号 12pt", br.font.size == Pt(12.0))
ind = body_p._p.pPr.ind
check("首行缩进 2 字符", ind is not None and ind.get(qn("w:firstLineChars")) == "200")
check("行距 1.5", body_p.paragraph_format.line_spacing == 1.5)
check("段后 6pt", body_p.paragraph_format.space_after == Pt(6.0))

# ---------- 2) 构造范本 → 解析 ----------
tpl = os.path.join(d, "tpl.docx")
td = Document()
tp = td.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("范本标题")
tr.font.name = "黑体"
tr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
tr.font.size = Pt(18.0)
tr.font.bold = True
for i, body_text in enumerate(["正文第一行内容。", "正文第二行内容。"]):
    bp = td.add_paragraph()
    bp.paragraph_format.line_spacing = 2.0
    bp.paragraph_format.space_after = Pt(8)
    bpPr = bp._p.get_or_add_pPr()
    ind = bpPr.get_or_add_ind()
    ind.set(qn("w:firstLineChars"), "200")
    ind.set(qn("w:firstLine"), "480")
    r = bp.add_run(body_text)
    r.font.name = "仿宋"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    r.font.size = Pt(14.0)
td.save(tpl)

parsed = parse_template(tpl)
check("范本标题字体", parsed.title_font == "黑体")
check("范本标题字号 18", parsed.title_size == 18.0)
check("范本标题加粗", parsed.title_bold is True)
check("范本标题居中", parsed.title_align == "center")
check("范本正文字体", parsed.body_font == "仿宋")
check("范本正文字号 14", parsed.body_size == 14.0)
check("范本首行缩进 2 字符", parsed.first_indent_chars == 2.0)
check("范本行距 2.0", parsed.line_spacing == 2.0)
check("范本段后 8", parsed.space_after == 8.0)
check("范本 describe 可读", "仿宋" in parsed.describe() and "居中" in parsed.describe())

# ---------- 3) 格式设置弹窗 ----------
dlg = ExportFormatDialog()
dlg.preset_combo.setCurrentText("论文")
f = dlg.fmt()
check("预设论文→字体/行距", f.body_font == "宋体" and f.line_spacing == 1.5 and f.first_indent_chars == 2.0)
dlg.preset_combo.setCurrentText("散文")
f = dlg.fmt()
check("预设散文→楷体", f.body_font == "楷体" and f.title_size == 18.0)
dlg.preset_combo.setCurrentText("自定义")
dlg.remember_check.setChecked(True)
check("remember 可读", dlg.remember() is True)
cfg_round = DocFormat.from_config(f.to_config())
check("DocFormat 配置往返", cfg_round.to_config() == f.to_config())
dlg.deleteLater()

# ---------- 4) 主窗口导出流程（记住设置 → 直接导出） ----------
win = MainWindow()
win.show()
win.config.setdefault("export", {})
win.config["export"]["docx_format_remembered"] = True
win.config["export"]["docx_format"] = PRESETS["默认"].to_config()

from app.models import Book, Chapter
from app.storage import Storage
dd = tempfile.mkdtemp()
bk = Book(title="散文集", author="A", book_type="散文随笔")
st = Storage.create_project(bk, dd)
c = Chapter(book_id=bk.id, title="雨夜", content="　　雨声敲打窗棂。")
c.id = st.add_chapter(c)
win._set_project(st)
win.open_chapter(c.id)
ed = win.current_editor()

out2 = os.path.join(dd, "exported.docx")
old_save = QFileDialog.getSaveFileName
QFileDialog.getSaveFileName = staticmethod(lambda *a, **k: (out2, "docx"))
try:
    ok_flag = win._export_docx_with_format(out2, ed.content(), "雨夜")
finally:
    QFileDialog.getSaveFileName = old_save
check("记住设置直接导出成功", ok_flag is True and os.path.exists(out2))
doc2 = Document(out2)
t2 = [p.text for p in doc2.paragraphs if p.text.strip()]
check("导出含标题雨夜", t2 and t2[0] == "雨夜")

# 按范本导出 plain 流程（patch QFileDialog 返回路径）
out3 = os.path.join(dd, "tpl_out.docx")
QFileDialog.getSaveFileName = staticmethod(lambda *a, **k: (out3, "docx"))
fired = []
try:
    win._tpl_save("标题行\n正文第一段。\n\n正文第二段。", PRESETS["默认"], fired.append, title=None, kind="docx")
finally:
    QFileDialog.getSaveFileName = old_save
check("plain 按范本导出完成", fired == [None] and os.path.exists(out3))
doc3 = Document(out3)
t3 = [p.text for p in doc3.paragraphs if p.text.strip()]
check("plain 导出标题取首行", t3 and t3[0] == "标题行")

# PDF 导出（格式化）：生成文件存在
out4 = os.path.join(dd, "out.pdf")
from app.exporter import export_pdf_formatted
export_pdf_formatted(out4, "　　第一段。\n　　第二段。", title="PDF标题", fmt=PRESETS["默认"])
check("PDF 导出生成文件", os.path.exists(out4) and os.path.getsize(out4) > 1000)

# 手动格式弹窗：current_fmt 构建正确
from app.dialogs.template_export_dialog import TemplateExportDialog
td = TemplateExportDialog()
td.manual_radio.setChecked(True)
mf = td.current_fmt()
check("手动格式行距可设", mf is not None and mf.line_spacing in (1.0, 1.15, 1.5, 2.0, 2.5, 3.0))
check("手动格式正文字体", mf.body_font in FONTS)
check("手动格式缩进 2 字符", mf.first_indent_chars == 2.0)
td.tpl_radio.setChecked(True)
check("范本模式未选范本 fmt=None", td.current_fmt() is None)

# 单选钮互斥修复：来源与模式分属不同组，模式可正常点选
check("默认范本模式选中", td.tpl_radio.isChecked())
td.mode_radios[1][1].setChecked(True)
check("模式单选钮可点选（互斥修复）", td.mode_radios[1][1].isChecked())
check("点模式不影响来源", td.tpl_radio.isChecked())
td.manual_radio.setChecked(True)
check("切来源后模式保持选中", td.mode_radios[1][1].isChecked())
# 输入框提示随模式变化
td.mode_radios[2][1].setChecked(True)
check("直接排版模式提示粘贴", "粘贴" in td.input_edit.placeholderText())
td.mode_radios[0][1].setChecked(True)
check("AI 生成模式提示写作要求", "写作要求" in td.input_edit.placeholderText())
td.deleteLater()

# ---------- 列表 + 页眉页脚 一比一 ----------
# 1) 导出含列表的文档：- 与 1. 开头行 → Word 原生列表样式
out5 = os.path.join(d, "list.docx")
export_docx_formatted(
    out5,
    "正文段落。\n- 无序项甲\n- 无序项乙\n1. 有序项一\n2. 有序项二\n结尾段落。",
    title="列表文档",
    fmt=DocFormat(name="列表"),
)
doc5 = Document(out5)
styles5 = [p.style.name for p in doc5.paragraphs if p.text.strip()]
check("无序列表用 List Bullet", any("List Bullet" in s for s in styles5))
check("有序列表用 List Number", any("List Number" in s for s in styles5))
check("普通段落保留", any(s not in ("List Bullet", "List Number") and "List" not in s for s in styles5))

# 2) 页眉页脚导出
out6 = os.path.join(d, "hf.docx")
hf_fmt = DocFormat(header_text="《山居笔记》", footer_text="第 1 页", name="页眉页脚")
export_docx_formatted(out6, "正文。", title="文章", fmt=hf_fmt)
doc6 = Document(out6)
hdr = "".join(p.text for p in doc6.sections[0].header.paragraphs)
ftr = "".join(p.text for p in doc6.sections[0].footer.paragraphs)
check("页眉导出", "《山居笔记》" in hdr)
check("页脚导出", "第 1 页" in ftr)

# 3) 构造含列表+页眉页脚的范本 → 解析
tpl2 = os.path.join(d, "tpl2.docx")
td2 = Document()
td2.sections[0].header.paragraphs[0].add_run("学校文件")
td2.sections[0].footer.paragraphs[0].add_run("机密")
tp = td2.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("文件标题")
tr.font.bold = True
tr.font.size = Pt(16.0)
td2.add_paragraph("正文第一段。")
td2.add_paragraph("要点一", style="List Bullet")
td2.add_paragraph("要点二", style="List Bullet")
td2.add_paragraph("步骤一", style="List Number")
td2.save(tpl2)
parsed2 = parse_template(tpl2)
check("范本解析出无序列表", parsed2.list_bullet is True)
check("范本解析出有序列表", parsed2.list_numbered is True)
check("范本解析出页眉", parsed2.header_text == "学校文件")
check("范本解析出页脚", parsed2.footer_text == "机密")
check("layout_instructions 含列表说明", "列表" in parsed2.layout_instructions())
check("layout_instructions 含页眉", "页眉" in parsed2.layout_instructions())
check("describe 含列表", "列表" in parsed2.describe())

# ---------- 漫剧脚本导出 ----------
from app.exporter import export_manju
manju = os.path.join(d, "manju.txt")
export_manju(manju, "　　林晚推开门，屋外风雪正紧。\n"
                     "　　林晚说：“我一定要找到师父！”\n"
                     "　　剑光一闪。\n"
                     "　　“师父？”她低声问。", "风雪夜")
mt = open(manju, encoding="utf-8").read()
check("漫剧导出含镜头", "【镜头 1】" in mt and "【镜头 2】" in mt)
check("漫剧导出含画面", "画面：" in mt)
check("漫剧台词提取说话人", "台词：林晚：" in mt)
check("漫剧导出标题", "《风雪夜》" in mt)

# 菜单入口
from app.main_window import MainWindow as _MW2
win_m = _MW2()
flat_m = []
for _name, menu in getattr(win_m, "_menus", []):
    for a in menu.actions():
        if not a.menu():
            flat_m.append(a.text())
check("菜单含导出漫剧", any("漫剧" in t for t in flat_m))
win_m.close()

# 打印：patch QPrintDialog.exec → 取消，流程不崩
from PySide6.QtPrintSupport import QPrintDialog
old_exec = QPrintDialog.exec
QPrintDialog.exec = lambda self: QPrintDialog.DialogCode.Rejected
try:
    win.print_current_chapter()
    check("打印取消不报错", True)
finally:
    QPrintDialog.exec = old_exec

# 菜单入口存在（主窗口菜单存于 self._menus：[(name, QMenu), ...]）
flat = []
for _name, menu in getattr(win, "_menus", []):
    for a in menu.actions():
        if not a.menu():
            flat.append(a.text())
check("菜单含按范本导出", any("按范本导出" in t for t in flat))
check("菜单含导出为Word", any("导出为 Word" in t for t in flat))
check("菜单含导出为PDF", any("导出为 PDF" in t for t in flat))
check("菜单含打印", any("打印" in t for t in flat))

win.close()
app.processEvents()

print("\nRESULT:", "ALL PASS" if ok else "HAS FAILURES")
sys.exit(0 if ok else 1)
